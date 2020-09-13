#!/usr/bin/python3

import requests
import docx
import lxml.html as html
import os
import datetime
import time


XPATH_LINK_TO_ARTICLE = '//text-fill[not(@class)]/a/@href'
XPATH_TITLE = '//div[@class="mb-auto"]/text-fill[not(@class)]/a/text()'
XPATH_SUMMARY = '//div[@class="lead"]/p/text()'
XPATH_BODY = '//div[@class="html-content"]/p[not(@class)]/text()'

HOME_URL = 'https://www.larepublica.co'


def parse_notice(link, today):
    ''' Extract links content and create word documents '''
    start = time.time()
    try:
        response = requests.get(link)
        if response.status_code == 200:
            notice = response.content.decode('utf-8')
            parsed = html.fromstring(notice)
            try:
                title = parsed.xpath(XPATH_TITLE)[0]
                title = title.replace('\"', '')
                summary = parsed.xpath(XPATH_SUMMARY)[0]
                body = parsed.xpath(XPATH_BODY)
            except IndexError:
                return

            doc = docx.Document()
            doc.add_paragraph(title)
            doc.add_paragraph(summary)
            for p in body:
                doc.add_paragraph(p)
            doc.save('{}/{}.docx'.format(today, title))
        else:
            raise ValueError('Error: {}'.format(response.status_code))
    except ValueError as ve:
        print(ve)
    end = time.time()
    print('Time: {}'.format(end - start))


def parse_home():
    ''' Extract all links '''
    try:
        response = requests.get(HOME_URL)
        if response.status_code == 200:
            home = response.content.decode('utf-8')
            parsed = html.fromstring(home)
            links_to_notices = parsed.xpath(XPATH_LINK_TO_ARTICLE)
            today = datetime.date.today().strftime('%d-%m-%Y')
            if not os.path.isdir(today):
                os.mkdir(today)

            for link in links_to_notices:
                parse_notice(link, today)
        else:
            raise ValueError('Error: {}'.format(response.status_code))
    except ValueError as ve:
        print(ve)


def run():
    start = time.time()
    local_time = time.ctime(start)
    print('Start: {}'.format(local_time))
    parse_home()
    end = time.time()
    local_time = time.ctime(end)
    print('End: {}'.format(local_time))
    total_secs = end - start
    print('Total time: {}'.format(total_secs))


if __name__ == "__main__":
    run()
